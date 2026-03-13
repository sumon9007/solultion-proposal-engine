#!/usr/bin/env python3
"""
export_word.py — Export an approved proposal to a branded Word document (.docx).

Usage:
    python scripts/export_word.py <path-to-approved-file>

The script:
1. Validates the file is an _approved.md
2. Parses markdown into python-docx structure
3. Applies brand colours and fonts from .env BRAND_* settings
4. Adds page header (company name / logo) and footer (page number, confidential)
5. Saves to output/exports/ with versioned filename

Dependencies:
    pip install python-docx

Returns exit code 0 on success, 1 on failure.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from helpers import (
    PROJECT_ROOT,
    approved_to_export_name_for,
    bold as bold_text,
    ensure_output_dirs,
    fail,
    info,
    load_env,
    next_export_version_for,
    ok,
    read_file,
    resolve_proposal_path,
    warn,
)


# ─── Colour Utilities ─────────────────────────────────────────────────────────

def _hex_to_rgb(hex_colour: str) -> tuple[int, int, int]:
    """Convert hex colour (#RRGGBB) to (r, g, b) int tuple."""
    h = hex_colour.lstrip("#")
    if len(h) != 6:
        raise ValueError(f"Invalid hex colour: #{h}")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


# ─── Document Styles ──────────────────────────────────────────────────────────

def _apply_brand_styles(doc, env: dict[str, str]) -> None:
    """Apply brand colours and fonts to the document's built-in paragraph styles."""
    from docx.shared import RGBColor, Pt

    primary_hex = env.get("BRAND_PRIMARY_COLOUR", "#0B5FFF")
    secondary_hex = env.get("BRAND_SECONDARY_COLOUR", "#0F172A")
    font_heading = env.get("BRAND_FONT_HEADINGS", "Aptos Display")
    font_body = env.get("BRAND_FONT_BODY", "Aptos")

    try:
        rp, gp, bp = _hex_to_rgb(primary_hex)
        rs, gs, bs = _hex_to_rgb(secondary_hex)
    except ValueError:
        rp, gp, bp = 11, 95, 255
        rs, gs, bs = 15, 23, 42

    style_map = {
        "Heading 1": (font_heading, Pt(20), RGBColor(rs, gs, bs), True),
        "Heading 2": (font_heading, Pt(14), RGBColor(rs, gs, bs), True),
        "Heading 3": (font_heading, Pt(12), RGBColor(rp, gp, bp), True),
        "Heading 4": (font_heading, Pt(11), RGBColor(rp, gp, bp), False),
    }

    for style_name, (font, size, colour, is_bold) in style_map.items():
        try:
            s = doc.styles[style_name]
            s.font.name = font
            s.font.size = size
            s.font.color.rgb = colour
            s.font.bold = is_bold
        except KeyError:
            pass

    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_body
        normal.font.size = Pt(11)
    except KeyError:
        pass


def _add_header_footer(doc, env: dict[str, str]) -> None:
    """Add running header (company name) and footer (page number | confidential)."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    company_name = env.get("COMPANY_NAME", "")
    secondary_hex = env.get("BRAND_SECONDARY_COLOUR", "#0F172A")
    try:
        rs, gs, bs = _hex_to_rgb(secondary_hex)
    except ValueError:
        rs, gs, bs = 15, 23, 42

    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.clear()
    run = h_para.add_run(company_name)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(rs, gs, bs)

    # Footer: page number | CONFIDENTIAL
    footer = section.footer
    footer.is_linked_to_previous = False
    f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    f_para.clear()
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run_pg = f_para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run_pg._r.append(fld)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_pg._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_pg._r.append(fld_end)

    run_conf = f_para.add_run("  |  CONFIDENTIAL")
    run_conf.font.size = Pt(8)
    run_conf.font.color.rgb = RGBColor(120, 136, 153)


# ─── Inline Markdown Parser ───────────────────────────────────────────────────

def _parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Parse inline markdown (bold, italic) into (text, is_bold, is_italic) tuples."""
    result: list[tuple[str, bool, bool]] = []
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            result.append((text[last_end:m.start()], False, False))
        if m.group(1):
            result.append((m.group(1), True, False))
        elif m.group(2):
            result.append((m.group(2), False, True))
        elif m.group(3):
            result.append((m.group(3), False, False))
        last_end = m.end()
    if last_end < len(text):
        result.append((text[last_end:], False, False))
    return result if result else [(text, False, False)]


def _add_para(doc, text: str, style: str = "Normal") -> None:
    """Add a paragraph with inline bold/italic formatting."""
    para = doc.add_paragraph(style=style)
    for fragment, is_bold, is_italic in _parse_inline(text):
        run = para.add_run(fragment)
        run.bold = is_bold
        run.italic = is_italic


# ─── Table Parser ─────────────────────────────────────────────────────────────

def _parse_md_table(table_lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into rows of cell strings, skipping separator rows."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c.replace(" ", "")) for c in cells if c):
            continue  # separator row
        rows.append(cells)
    return rows


def _add_table(doc, rows: list[list[str]], env: dict[str, str]) -> None:
    """Add a Word table with branded header row."""
    from docx.shared import RGBColor, Pt

    primary_hex = env.get("BRAND_PRIMARY_COLOUR", "#0B5FFF")
    secondary_hex = env.get("BRAND_SECONDARY_COLOUR", "#0F172A")
    try:
        rp, gp, bp = _hex_to_rgb(primary_hex)
        rs, gs, bs = _hex_to_rgb(secondary_hex)
    except ValueError:
        rp, gp, bp = 11, 95, 255
        rs, gs, bs = 15, 23, 42

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx in range(n_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            cell = table.rows[row_idx].cells[col_idx]
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(cell_text)
            if row_idx == 0:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(rs, gs, bs)


# ─── Markdown → Word Converter ────────────────────────────────────────────────

def convert_to_word(md_content: str, env: dict[str, str]):
    """Convert markdown content to a branded python-docx Document object."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _apply_brand_styles(doc, env)
    _add_header_footer(doc, env)

    heading_style_map = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
    }

    lines = md_content.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style = heading_style_map.get(level, "Heading 4")
            _add_para(doc, text, style=style)
            i += 1
            continue

        # Horizontal rule — thin paragraph
        if re.match(r"^[-_*]{3,}\s*$", line):
            doc.add_paragraph("─" * 50)
            i += 1
            continue

        # Table — collect all contiguous table lines
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_md_table(table_lines)
            if rows:
                _add_table(doc, rows, env)
            continue

        # Unordered list
        m = re.match(r"^\s*[-*]\s+(.+)", line)
        if m:
            _add_para(doc, m.group(1), style="List Bullet")
            i += 1
            continue

        # Ordered list
        m = re.match(r"^\s*\d+\.\s+(.+)", line)
        if m:
            _add_para(doc, m.group(1), style="List Number")
            i += 1
            continue

        # Blockquote
        m = re.match(r"^>\s*(.*)", line)
        if m:
            _add_para(doc, m.group(1), style="Quote")
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph("\n".join(code_lines))
                for run in para.runs:
                    run.font.name = "Cascadia Mono"
                    run.font.size = Pt(9)
            i += 1  # skip closing ```
            continue

        # Regular paragraph
        _add_para(doc, line)
        i += 1

    return doc


# ─── Export ───────────────────────────────────────────────────────────────────

def export_word(approved_path: Path) -> bool:
    ensure_output_dirs()

    print()
    print(bold_text(f"EXPORTING TO WORD: {approved_path.name}"))
    print("=" * 60)

    # ── Guard: must be approved file ─────────────────────────────
    if "_approved.md" not in approved_path.name:
        if "_draft.md" in approved_path.name:
            print(fail("Cannot export a draft file. Approve it first:"))
            print(info(f"  python scripts/approve_draft.py {approved_path}"))
        else:
            print(fail("File does not appear to be an approved proposal."))
            print(warn("Expected filename pattern: *_approved.md"))
        return False

    # ── Check dependency ─────────────────────────────────────────
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print(fail("python-docx is not installed."))
        print(info("  pip install python-docx"))
        return False

    # ── Read content ─────────────────────────────────────────────
    md_content = read_file(approved_path)
    env = load_env()

    print(info("  Applying brand styles from .env..."))
    doc = convert_to_word(md_content, env)

    # ── Determine output path ─────────────────────────────────────
    version = next_export_version_for(approved_path, "docx")
    export_path = approved_to_export_name_for(approved_path, "docx", version)

    # ── Write output ──────────────────────────────────────────────
    doc.save(str(export_path))
    print(ok(f"Exported: {export_path}"))
    print(ok(f"Version:  v{version}"))
    print()
    print(bold_text("NEXT STEP:"))
    print(info(f"  Open document: {export_path.resolve()}"))
    print()

    return True


# ─── Entry Point ──────────────────────────────────────────────────────────────

def main() -> None:
    if len(sys.argv) < 2:
        print(f"Usage: python {Path(__file__).name} <path-to-approved-file>")
        sys.exit(1)

    path = resolve_proposal_path(sys.argv[1])

    if not path.exists():
        print(fail(f"File not found: {path}"))
        sys.exit(1)

    success = export_word(path)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
